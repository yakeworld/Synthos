#!/usr/bin/env python3
"""
Generic helper for filling DOCX form tables.

Usage:
    python3 fill_form.py --input template.docx --output filled.docx --field name "Project Name" --field description "Long description here..."
    
Or programmatically:
    from fill_form import fill_table_cell, set_cell_content
    fill_table_cell(doc, 0, 1, "Project Name", font_bold=True, font_size=14)
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse
import sys


def ensure_rPr(run):
    """Ensure run has rPr element."""
    if run._element.rPr is None:
        run._element.insert(0, OxmlElement('w:rPr'))


def set_chinese_font(run, font_name="宋体"):
    """Set Chinese font for a run."""
    ensure_rPr(run)
    rPr = run._element.rPr
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def fill_table_cell(doc, row_idx, cell_idx, text, font_size=10.5, bold=False, 
                    font_name="宋体", alignment=None, color=None):
    """Fill a table cell with text and formatting."""
    table = doc.tables[0]  # Assumes first table; modify if needed
    cell = table.rows[row_idx].cells[cell_idx]
    
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    set_chinese_font(run, font_name)
    run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    
    return run


def fill_multi_cell(doc, row_idx, start_cell, text, font_size=10.5, bold=False, 
                    font_name="宋体", num_cells=3):
    """Fill content across multiple cells (for merged content areas)."""
    for i in range(num_cells):
        cell_idx = start_cell + i
        cell = doc.tables[0].rows[row_idx].cells[cell_idx]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        set_chinese_font(run, font_name)
        run.font.size = Pt(font_size)
        if bold:
            run.font.bold = True
    return doc


def set_name_cell(doc, name, font_size=14):
    """Set the smart agent name in row 0."""
    fill_table_cell(doc, 0, 1, name, font_size=font_size, bold=True, font_name="黑体")


def set_content_cells(doc, row_idx, content, font_size=10.5):
    """Set content across cells 1-3 of a row."""
    fill_multi_cell(doc, row_idx, 1, content, font_size=font_size, num_cells=3)


def set_checkbox_row(doc, row_idx, options_text, font_size=10):
    """Set checkbox options."""
    fill_table_cell(doc, row_idx, 1, options_text, font_size=font_size, font_name="宋体")


def set_link_cell(doc, row_idx, link_text, font_size=10.5, italic=True):
    """Set a URL or link."""
    fill_table_cell(doc, row_idx, 1, link_text, font_size=font_size, font_name="宋体", bold=False)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Fill DOCX form template")
    parser.add_argument("--input", required=True, help="Input template DOCX")
    parser.add_argument("--output", required=True, help="Output filled DOCX")
    args = parser.parse_args()
    
    doc = Document(args.input)
    
    # Example: fill a basic form
    print(f"Read {args.input}: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    doc.save(args.output)
    print(f"Saved {args.output}")
