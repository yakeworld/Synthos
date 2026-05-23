#!/usr/bin/env python3
"""
将 Markdown 转为 Word（.docx），按标题层级映射为 Word 内置「标题 1–9」样式，
便于交底书交付代理人或所内流程。

支持：ATX 标题 (#–######)、段落、**粗体**、行内 `代码`、无序/有序列表、
围栏代码块、简单 GFM 表格、引用块（>）、水平线（---）、行内图片 ``![](path.png)``
（在最大宽、最大高约束下**等比缩放**，竖图自动缩小宽度以整图落入版面）。

**连续多行正文**（中间无空行、且非列表/标题等）时，**每一行**输出为 Word 中**独立一段**，
以便「（1）…（2）…」等分条换行；若须在同一段内接排，请写**同一行**内或用 Markdown 空行分隔逻辑段。

定稿宜先用同目录 **`mermaid_render.py`** 将 **mermaid** 转为 PNG；若个别块生图失败仍保留 `` ```mermaid`` 围栏，本文档会将其作为**代码块**写入 Word。

用法：
  python md_to_docx.py --input disclosure.md --output disclosure.docx
  python md_to_docx.py -i a.md -o b.docx --base-dir .   # 解析图片相对路径

依赖：python-docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# 插图最大尺寸（英寸）：在常见 A4、默认边距下保证整图可见、按比例缩放（不过宽也不过高）。
_DEFAULT_IMAGE_MAX_W_IN = 5.5
_DEFAULT_IMAGE_MAX_H_IN = 8.2


def _image_pixel_size(path: Path) -> tuple[int, int] | None:
    """读取常见位图宽高（像素），失败返回 None。不依赖 Pillow。"""
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if len(raw) >= 24 and raw.startswith(b"\x89PNG\r\n\x1a\n") and raw[12:16] == b"IHDR":
        w = int.from_bytes(raw[16:20], "big")
        h = int.from_bytes(raw[20:24], "big")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 10 and raw[:3] == b"GIF" and raw[3:6] in (b"87a", b"89a"):
        w = int.from_bytes(raw[6:8], "little")
        h = int.from_bytes(raw[8:10], "little")
        if w > 0 and h > 0:
            return w, h
    if len(raw) >= 4 and raw.startswith(b"\xff\xd8"):
        i = 2
        n = len(raw)
        while i < n:
            if raw[i] != 0xFF:
                i += 1
                continue
            i += 1
            while i < n and raw[i] == 0xFF:
                i += 1
            if i >= n:
                break
            marker = raw[i]
            i += 1
            if marker in (0xD8, 0xD9):
                continue
            if marker == 0xDA:
                break
            if 0xD0 <= marker <= 0xD7:
                continue
            if i + 2 > n:
                break
            seg_len = int.from_bytes(raw[i : i + 2], "big")
            if seg_len < 2:
                break
            i += 2
            if marker in (0xC0, 0xC1, 0xC2) and i + 5 <= n:
                h = int.from_bytes(raw[i + 1 : i + 3], "big")
                w = int.from_bytes(raw[i + 3 : i + 5], "big")
                if w > 0 and h > 0:
                    return w, h
            i += seg_len - 2
    return None


def _fit_image_display_inches(
    px_w: int,
    px_h: int,
    *,
    max_w_in: float,
    max_h_in: float,
) -> tuple[Inches, Inches]:
    """在不超过 max_w / max_h 的前提下等比缩放，使整图落入版面。"""
    if px_w <= 0 or px_h <= 0:
        return Inches(max_w_in), Inches(max_h_in * 0.5)
    aw = max_w_in
    ah = aw * px_h / px_w
    if ah > max_h_in:
        ah = max_h_in
        aw = ah * px_w / px_h
    return Inches(aw), Inches(ah)


def _set_run_font(run, name: str = "宋体", size_pt: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def _add_inline_to_paragraph(paragraph, text: str, *, mono: bool = False):
    """解析 **粗体**、`行内代码` 与普通文本，写入同一段落。"""
    if not text:
        return
    # 拆分为：粗体、行内代码、普通
    pattern = re.compile(r"(\*\*[^*]+?\*\*|`[^`]+?`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            _set_run_font(run, "Consolas" if mono else "宋体", 10.5 if not mono else 9)
        token = m.group(1)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, "宋体", 10.5, bold=True)
        else:  # `code`
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, "Consolas", 9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, "Consolas" if mono else "宋体", 10.5 if not mono else 9)


def _add_heading(doc: Document, level: int, text: str):
    """level 1–9 对应 Word 标题 1–标题 9；去除行内标记时保留可读文本。"""
    plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    plain = re.sub(r"`([^`]+)`", r"\1", plain)
    h = doc.add_heading(plain, level=min(max(level, 1), 9))
    for run in h.runs:
        _set_run_font(run, "黑体" if level <= 2 else "宋体")


def _add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    _add_inline_to_paragraph(p, text)
    for run in p.runs:
        if run.font.name in (None, ""):
            _set_run_font(run, "宋体", 10.5)


def _add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    body = "\n".join(lines)
    run = p.add_run(body)
    _set_run_font(run, "Consolas", 9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def _add_list_item(doc: Document, text: str, ordered: bool, base_dir: Path | None):
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except (KeyError, ValueError):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(3)
    _add_inline_to_paragraph(p, text)
    for run in p.runs:
        _set_run_font(run, "宋体", 10.5)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def _parse_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_table_sep(row: list[str]) -> bool:
    if not row:
        return False
    return all(re.match(r"^:?-{3,}:?$", c.strip()) for c in row if c.strip())


def _add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_to_paragraph(p, cell_text)
            for run in p.runs:
                _set_run_font(run, "宋体", 10)


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("─" * 32)
    _set_run_font(run, "宋体", 8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _try_add_image(
    doc: Document,
    line: str,
    base_dir: Path | None,
    *,
    max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> bool:
    m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not m or not base_dir:
        return False
    alt, src = m.group(1), m.group(2).strip()
    path = (base_dir / src).resolve() if not Path(src).is_absolute() else Path(src)
    if not path.is_file():
        p = doc.add_paragraph()
        p.add_run(f"[图片缺失: {alt or src} — {path}]")
        return True
    try:
        dims = _image_pixel_size(path)
        if dims:
            w_in, h_in = _fit_image_display_inches(
                *dims, max_w_in=max_w_in, max_h_in=max_h_in
            )
            doc.add_picture(str(path), width=w_in, height=h_in)
        else:
            doc.add_picture(str(path), width=Inches(max_w_in))
    except Exception:
        p = doc.add_paragraph()
        p.add_run(f"[图片无法嵌入: {path}]")
    return True


def convert_md_to_docx(
    md_text: str,
    base_dir: Path | None,
    *,
    image_max_w_in: float = _DEFAULT_IMAGE_MAX_W_IN,
    image_max_h_in: float = _DEFAULT_IMAGE_MAX_H_IN,
) -> Document:
    doc = Document()
    # 默认正文样式
    try:
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        if style._element.rPr is not None:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
    except (AttributeError, KeyError):
        pass

    lines = md_text.splitlines()
    i = 0
    para_buf: list[str] = []

    def flush_paragraph():
        nonlocal para_buf
        if not para_buf:
            return
        # 每行独立成段，避免「（1）…\n（2）…」被空格拼成一段（Word 内不换行）
        for p in para_buf:
            t = p.strip()
            if t:
                _add_body_paragraph(doc, t)
        para_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")

        if line.strip() == "":
            flush_paragraph()
            i += 1
            continue

        # 围栏代码块
        if line.strip().startswith("```"):
            flush_paragraph()
            fence_lang = line.strip()[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            _add_code_block(doc, code_lines)
            continue

        # 图片独占一行
        if line.strip().startswith("![") and "](" in line:
            flush_paragraph()
            _try_add_image(
                doc,
                line,
                base_dir,
                max_w_in=image_max_w_in,
                max_h_in=image_max_h_in,
            )
            i += 1
            continue

        # 水平线
        if re.match(r"^[\s\-*_]{3,}\s*$", line) and set(line.strip()) <= {"-", "*", "_", " "}:
            flush_paragraph()
            _add_horizontal_rule(doc)
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            flush_paragraph()
            level = len(m.group(1))
            title = m.group(2).strip()
            title = re.sub(r"\s+#+\s*$", "", title)
            _add_heading(doc, level, title)
            i += 1
            continue

        # 引用
        if line.lstrip().startswith("> "):
            flush_paragraph()
            quote = line.lstrip()[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            _add_inline_to_paragraph(p, quote)
            for run in p.runs:
                _set_run_font(run, "宋体", 10.5)
            i += 1
            continue

        # 表格块
        if _is_table_row(line):
            flush_paragraph()
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i]):
                row = _parse_table_row(lines[i])
                if not _is_table_sep(row):
                    table_rows.append(row)
                i += 1
            _add_table(doc, table_rows)
            continue

        # 无序列表
        um = re.match(r"^(\s*)[-*+]\s+(.+)$", line)
        if um:
            flush_paragraph()
            _add_list_item(doc, um.group(2).strip(), ordered=False, base_dir=base_dir)
            i += 1
            continue

        # 有序列表
        om = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if om:
            flush_paragraph()
            _add_list_item(doc, om.group(2).strip(), ordered=True, base_dir=base_dir)
            i += 1
            continue

        para_buf.append(line)
        i += 1

    flush_paragraph()
    return doc


def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description="Markdown → Word（标题样式映射）")
    p.add_argument("-i", "--input", required=True, help="输入 .md 路径")
    p.add_argument("-o", "--output", required=True, help="输出 .docx 路径")
    p.add_argument(
        "--base-dir",
        default=None,
        help="解析 ![](/相对路径) 图片时的根目录（默认使用 .md 所在目录）",
    )
    p.add_argument(
        "--image-max-width-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_W_IN,
        metavar="IN",
        help=f"插图最大宽度（英寸，默认 {_DEFAULT_IMAGE_MAX_W_IN}），与高度共同约束等比缩放",
    )
    p.add_argument(
        "--image-max-height-inches",
        type=float,
        default=_DEFAULT_IMAGE_MAX_H_IN,
        metavar="IN",
        help=f"插图最大高度（英寸，默认 {_DEFAULT_IMAGE_MAX_H_IN}），避免竖图仅按宽度缩放后超出单页可视区域",
    )
    args = p.parse_args(argv)

    in_path = Path(args.input).resolve()
    if not in_path.is_file():
        print(f"错误：找不到输入文件 {in_path}", file=sys.stderr)
        return 1

    base = Path(args.base_dir).resolve() if args.base_dir else in_path.parent
    try:
        md_text = in_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        md_text = in_path.read_text(encoding="utf-8", errors="replace")
        print("警告：输入文件含非 UTF-8 字节，已使用替换字符解码后继续转换。", file=sys.stderr)

    doc = convert_md_to_docx(
        md_text,
        base_dir=base,
        image_max_w_in=args.image_max_width_inches,
        image_max_h_in=args.image_max_height_inches,
    )
    out_path = Path(args.output).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"已写入: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
