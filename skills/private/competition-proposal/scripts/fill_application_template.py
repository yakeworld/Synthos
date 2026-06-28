#!/usr/bin/env python3
"""
Template: DOCX competition form filler.
1. Edit application_content.json with your content
2. Set SRC to your template path
3. Run: python3 fill_application.py
"""
import json, os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SRC = 'template.docx'       # Change to your template path
DST = 'filled_application.docx'  # Output path
CONTENT_FILE = 'application_content.json'

def set_cell(table, row_idx, col_idx, text):
    cell = table.rows[row_idx].cells[col_idx]
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.text = text
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Load content
with open(CONTENT_FILE, 'r', encoding='utf-8') as f:
    texts = json.load(f)

doc = Document(SRC)
table = doc.tables[0]

# === TABLE FIELDS ===
set_cell(table, 1, 2, texts.get('project_name', ''))

# === LONG TEXT SECTIONS ===
# Adjust paragraph indices to match your template
para_map = {
    31: 'project_bg',
    # ... add more mappings
}
for para_idx, key in para_map.items():
    if para_idx < len(doc.paragraphs) and key in texts:
        p = doc.paragraphs[para_idx]
        for r in p.runs:
            r.text = ''
        p.text = texts[key]

doc.save(DST)
print(f'Saved to {DST}')
